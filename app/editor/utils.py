# editor/utils.py

import os
from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Define heading formats
heading_formats = {
    1: {'bold': True, 'italic': False, 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
    2: {'bold': True, 'italic': False, 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
    3: {'bold': True, 'italic': True, 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
    4: {'bold': False, 'italic': True, 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
    5: {'bold': False, 'italic': True, 'alignment': WD_ALIGN_PARAGRAPH.LEFT}
}

def format_document(doc):
    # 1. Font: Set the entire document's font to Times New Roman.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'

    # 2. Font Size: Set the font size to 14.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(14)

    # 3. Document Size: Ensure the document is set to A4 size.
    for section in doc.sections:
        section.page_height = Cm(29.7)  # A4 height
        section.page_width = Cm(21)  # A4 width

    # 4. Margins: Set the margins to 2 cm top and bottom, 3 cm left, 2 cm right.
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # 5. Text Alignment: Justify all text.
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 6. Line spacing: 1.5
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5

    # 7. Paragraph first line indent: 1.27cm
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.first_line_indent = Cm(1.27)

    return doc

def set_normal_style(doc):
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(14)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return doc

def set_heading_styles(doc):
    styles = doc.styles
    for level in heading_formats:
        style_name = f'Heading {level}'
        try:
            style = styles[style_name]
        except KeyError:
            # Style doesn't exist, create it
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        style.quick_style = True
        style.font.color.rgb = None
        style.font.bold = heading_formats[level]['bold']
        style.font.italic = heading_formats[level]['italic']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(13)
        style.paragraph_format.alignment = heading_formats[level]['alignment']
        style.paragraph_format.line_spacing = 1.5  # Set line spacing to 1.5
        if level in [1, 2]:
            style.paragraph_format.first_line_indent = Cm(0)  # No indent for Heading 1 and 2
        else:
            style.paragraph_format.first_line_indent = Cm(1.27)  # Indent for other headings
    return doc

def get_paragraphs_and_headings(document):
    paragraphs = []
    headings = []
    for i, para in enumerate(document.paragraphs):
        text = para.text
        style = para.style.name
        paragraphs.append({'index': i, 'text': text, 'style': style})

        # Check if the paragraph is a heading
        if style.startswith('Heading '):
            try:
                level = int(style.replace('Heading ', ''))
            except ValueError:
                level = 1  # Default to level 1 if parsing fails
            indent = (level - 1) * 20  # Calculate indentation
            headings.append({'index': i, 'text': text, 'level': level, 'indent': indent})
    return paragraphs, headings